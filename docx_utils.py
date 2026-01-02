# docx_utils.py
# -*- coding: utf-8 -*-
import re
import tempfile
from io import BytesIO
from zipfile import ZipFile
from pathlib import Path

import requests
import qrcode
from docx import Document


def generar_qr_y_barcode(url_qr: str, rfc: str):
    # --- QR ---
    qr = qrcode.QRCode(
        version=None,
        box_size=8,
        border=2,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
    )
    qr.add_data(url_qr)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white")

    buf_qr = BytesIO()
    qr_img.save(buf_qr, format="PNG")
    qr_bytes = buf_qr.getvalue()

    # --- Código de barras (servicio externo) ---
    import urllib.parse
    rfc_encoded = urllib.parse.quote_plus(rfc)

    url_barcode = (
        "https://barcode.tec-it.com/barcode.ashx"
        f"?data={rfc_encoded}"
        "&code=Code128"
        "&translate-esc=on"
        "&dpi=300"
    )
    resp = requests.get(url_barcode, timeout=20)
    resp.raise_for_status()
    barcode_bytes = resp.content

    return qr_bytes, barcode_bytes


def reemplazar_en_documento(ruta_entrada: str, ruta_salida: str, datos: dict):
    """
    Reemplaza placeholders tipo {{ RFC }} en document.xml / header / footer,
    y sustituye las imágenes word/media/image2.png (QR) y image6.png (barcode).
    Luego repara placeholders partidos en runs con python-docx.
    """
    rfc_val = datos.get("RFC_ETIQUETA") or datos.get("RFC", "")
    idcif_val = datos.get("IDCIF_ETIQUETA") or datos.get("IDCIF") or ""

    d3 = f"{idcif_val}_{rfc_val}" if idcif_val else ""
    url_qr = f"https://siat.sat.validacion-sat.com/app/qr/faces/pages/mobile/validadorqr.jsf?D1=10&D2=1&D3={d3}"

    qr_bytes, barcode_bytes = generar_qr_y_barcode(url_qr, rfc_val)

    placeholders = {
        "{{ RFC ETIQUETA }}": datos.get("RFC_ETIQUETA", ""),
        "{{ NOMBRE ETIQUETA }}": datos.get("NOMBRE_ETIQUETA", ""),
        "{{ idCIF }}": datos.get("IDCIF_ETIQUETA", datos.get("IDCIF", "")),
        "{{ FECHA }}": datos.get("FECHA", ""),
        "{{ FECHA CORTA }}": datos.get("FECHA_CORTA", ""),
        "{{ RFC }}": datos.get("RFC", ""),
        "{{ CURP }}": datos.get("CURP", ""),
        "{{ NOMBRE }}": datos.get("NOMBRE", ""),
        "{{ PRIMER APELLIDO }}": datos.get("PRIMER_APELLIDO", ""),
        "{{ SEGUNDO APELLIDO }}": datos.get("SEGUNDO_APELLIDO", ""),
        "{{ FECHA INICIO }}": datos.get("FECHA_INICIO", ""),
        "{{ ESTATUS }}": datos.get("ESTATUS", ""),
        "{{ FECHA ULTIMO }}": datos.get("FECHA_ULTIMO", ""),
        "{{ CP }}": datos.get("CP", ""),
        "{{ TIPO VIALIDAD }}": datos.get("TIPO_VIALIDAD", ""),
        "{{ VIALIDAD }}": datos.get("VIALIDAD", ""),
        "{{ NO EXTERIOR }}": datos.get("NO_EXTERIOR", ""),
        "{{ NO INTERIOR }}": datos.get("NO_INTERIOR", ""),
        "{{ COLONIA }}": datos.get("COLONIA", ""),
        "{{ LOCALIDAD }}": datos.get("LOCALIDAD", ""),
        "{{ ENTIDAD }}": datos.get("ENTIDAD", ""),
        "{{ REGIMEN }}": datos.get("REGIMEN", ""),
        "{{ FECHA ALTA }}": datos.get("FECHA_ALTA", ""),
    }

    # 1) Reemplazo directo en XML dentro del DOCX (zip)
    with ZipFile(ruta_entrada, "r") as zin, ZipFile(ruta_salida, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if (
                item.filename == "word/document.xml"
                or item.filename.startswith("word/header")
                or item.filename.startswith("word/footer")
            ):
                try:
                    xml_text = data.decode("utf-8")
                except UnicodeDecodeError:
                    pass
                else:
                    # Caso idCIF partido en runs: {{ idCIF }}
                    if idcif_val:
                        patron_idcif = r"<w:t>{{</w:t>.*?<w:t>idCIF</w:t>.*?<w:t>}}</w:t>"
                        xml_text, _ = re.subn(
                            patron_idcif,
                            f"<w:t>{idcif_val}</w:t>",
                            xml_text,
                            flags=re.DOTALL,
                        )

                    for k, v in placeholders.items():
                        if k in xml_text:
                            xml_text = xml_text.replace(k, v)

                    data = xml_text.encode("utf-8")

            # 2) Sustituir imágenes por nombre interno (ajusta si tu plantilla usa otros indices)
            if item.filename == "word/media/image2.png":
                data = qr_bytes
            elif item.filename == "word/media/image6.png":
                data = barcode_bytes

            zout.writestr(item, data)

    # 3) Reparación extra: placeholders divididos en runs dentro de párrafos/tablas
    doc = Document(ruta_salida)

    par_placeholders = {
        "{{ FECHA CORTA }}": datos.get("FECHA_CORTA", ""),
        "{{FECHA CORTA}}": datos.get("FECHA_CORTA", ""),
        "{{ FECHA }}": datos.get("FECHA", ""),
        "{{FECHA}}": datos.get("FECHA", ""),
        "{{ RFC }}": datos.get("RFC", ""),
        "{{RFC}}": datos.get("RFC", ""),
        "{{ idCIF }}": datos.get("IDCIF_ETIQUETA", datos.get("IDCIF", "")),
        "{{idCIF}}": datos.get("IDCIF_ETIQUETA", datos.get("IDCIF", "")),
    }

    def reemplazar_en_parrafos(paragraphs):
        for p in paragraphs:
            if "{{" not in p.text:
                continue

            full = "".join(r.text for r in p.runs)
            if "{{" not in full:
                continue

            start_idx = full.find("{{")
            if start_idx == -1:
                continue

            acc = 0
            start_run = None
            for i, r in enumerate(p.runs):
                if acc + len(r.text) > start_idx:
                    start_run = i
                    break
                acc += len(r.text)

            if start_run is None:
                continue

            suffix = "".join(r.text for r in p.runs[start_run:])
            new_suffix = suffix
            for k, v in par_placeholders.items():
                if k in new_suffix:
                    new_suffix = new_suffix.replace(k, v)

            if new_suffix == suffix:
                continue

            p.runs[start_run].text = new_suffix
            for r in p.runs[start_run + 1:]:
                r.text = ""

    reemplazar_en_parrafos(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_en_parrafos(cell.paragraphs)

    doc.save(ruta_salida)


def generar_docx_desde_plantilla(datos: dict, plantilla_path: str) -> str:
    """
    Genera un DOCX temporal desde una plantilla y devuelve la ruta al archivo final.
    """
    tmpdir = tempfile.mkdtemp()
    salida = str(Path(tmpdir) / "constancia.docx")
    reemplazar_en_documento(plantilla_path, salida, datos)
    return salida
